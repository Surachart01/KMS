import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_document(input_path, output_path):
    doc = Document(input_path)
    
    # We want font sizes:
    # Heading 1, 2 (maybe 2 is main as well? 3.X is H2) -> Let's treat H1 and H2 as main headings (18pt bold)
    # H3, H4 as subheadings (16pt bold)
    # Normal and List Paragraph as body (16pt normal, indented)
    
    black_color = RGBColor(0, 0, 0)
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        
        is_main_heading = style_name.startswith('Heading 1') or style_name.startswith('Heading 2')
        is_sub_heading = style_name.startswith('Heading 3') or style_name.startswith('Heading 4')
        is_heading = is_main_heading or is_sub_heading
        is_normal = not is_heading
        
        # Apply paragraph indentation for Normal body paragraphs
        if style_name == 'Normal':
            # Ignore paragraphs with image or just empty
            if para.text.strip() != '' and not para.text.startswith('ภาพที่') and not para.text.startswith('ตารางที่'):
                para.paragraph_format.first_line_indent = Inches(0.5)
            
            # center captions
            if para.text.startswith('ภาพที่') or para.text.startswith('ตารางที่'):
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        if style_name == 'Caption':
             para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for run in para.runs:
            run.font.color.rgb = black_color
            # Set font name to TH Sarabun PSK if possible, or just change size
            run.font.name = 'TH Sarabun PSK'
            
            if is_main_heading:
                run.font.size = Pt(18)
                run.font.bold = True
            elif is_sub_heading:
                run.font.size = Pt(16)
                run.font.bold = True
            else:
                run.font.size = Pt(16)
                # Keep bold if it was bold originally (e.g. bolded words in text)
                
    # Also fix tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = black_color
                        run.font.name = 'TH Sarabun PSK'
                        run.font.size = Pt(16)
    
    doc.save(output_path)

if __name__ == "__main__":
    format_document('ch3_temp.docx', 'ch3.docx')
    print("Formatting complete.")
