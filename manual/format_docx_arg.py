import sys
import subprocess

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Twips, Emu
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Twips, Emu
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy


def set_page_margins(doc):
    """Set A4 page with Thai academic margins: top/bottom 2.54cm, left 3cm, right 2.54cm"""
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.54)


def set_para_spacing(para, before=0, after=6, line_spacing=None):
    """Set paragraph spacing (in pt)"""
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    if line_spacing is not None:
        fmt.line_spacing = Pt(line_spacing)


def apply_run_style(run, size, bold=False, color=RGBColor(0, 0, 0)):
    run.font.name  = 'TH Sarabun PSK'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    # Force Thai font via XML
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'TH Sarabun PSK')
    rFonts.set(qn('w:hAnsi'), 'TH Sarabun PSK')
    rFonts.set(qn('w:cs'),    'TH Sarabun PSK')
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def add_page_break_before(para):
    """Add page break before a paragraph"""
    fmt = para.paragraph_format
    fmt.page_break_before = True


def is_appendix_heading(text):
    """Detect major appendix title lines like ภาคผนวก ก-1"""
    t = text.strip()
    return t.startswith('ภาคผนวก') and ('ก-' in t or 'ข-' in t or 'ค-' in t)


def style_table(table):
    """Apply consistent styling to a table"""
    black = RGBColor(0, 0, 0)
    # Table borders
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'),  'single')
        el.set(qn('w:sz'),   '4')
        el.set(qn('w:space'),'0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

    # Font & alignment in cells
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                fmt = para.paragraph_format
                fmt.space_before = Pt(2)
                fmt.space_after  = Pt(2)
                fmt.left_indent  = Inches(0.05)
                is_header_row = (i == 0)
                for run in para.runs:
                    apply_run_style(run, size=16,
                                    bold=is_header_row,
                                    color=black)
                # If header row, shade it slightly
                if is_header_row:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  'D9D9D9')  # light grey
                    existing_shd = tcPr.find(qn('w:shd'))
                    if existing_shd is not None:
                        tcPr.remove(existing_shd)
                    tcPr.append(shd)


def format_document(input_path, output_path):
    doc = Document(input_path)
    black = RGBColor(0, 0, 0)

    # 1. Page setup
    set_page_margins(doc)

    # 2. Paragraphs
    first_appendix = True
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        text = para.text.strip()

        is_h1 = style_name.startswith('Heading 1')
        is_h2 = style_name.startswith('Heading 2')
        is_h3 = style_name.startswith('Heading 3') or style_name.startswith('Heading 4')

        # --- Appendix title (Heading 1 that starts with ภาคผนวก) ---
        if is_h1 and is_appendix_heading(text):
            if first_appendix:
                first_appendix = False
            else:
                add_page_break_before(para)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_para_spacing(para, before=0, after=12, line_spacing=22)
            for run in para.runs:
                apply_run_style(run, size=18, bold=True, color=black)
            continue

        # --- Normal Heading 1 ---
        if is_h1:
            set_para_spacing(para, before=12, after=6, line_spacing=22)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in para.runs:
                apply_run_style(run, size=18, bold=True, color=black)
            continue

        # --- Heading 2 ---
        if is_h2:
            set_para_spacing(para, before=10, after=4, line_spacing=20)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in para.runs:
                apply_run_style(run, size=16, bold=True, color=black)
            continue

        # --- Heading 3 / 4 ---
        if is_h3:
            set_para_spacing(para, before=8, after=3, line_spacing=20)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in para.runs:
                apply_run_style(run, size=16, bold=True, color=black)
            continue

        # --- Caption / Figure / Table label ---
        if style_name == 'Caption' or text.startswith('ภาพที่') or text.startswith('ตารางที่'):
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_para_spacing(para, before=2, after=4, line_spacing=20)
            for run in para.runs:
                apply_run_style(run, size=16, bold=False, color=black)
            continue

        # --- List items (bullet / ordered) ---
        if 'List' in style_name:
            set_para_spacing(para, before=0, after=3, line_spacing=20)
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.first_line_indent = Pt(0)
            for run in para.runs:
                apply_run_style(run, size=16, bold=False, color=black)
            continue

        # --- Normal / body text ---
        if text != '':
            set_para_spacing(para, before=0, after=6, line_spacing=22)
            para.paragraph_format.first_line_indent = Inches(0.5)
            para.paragraph_format.left_indent = Pt(0)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            # Empty paragraph: minimal spacing
            set_para_spacing(para, before=0, after=0, line_spacing=20)

        for run in para.runs:
            apply_run_style(run, size=16, bold=False, color=black)

    # 3. Tables
    for table in doc.tables:
        style_table(table)

    # 4. Resize images to fit within text width and center them
    # Calculate usable text width in EMU
    section = doc.sections[0]
    text_width_emu = (section.page_width
                      - section.left_margin
                      - section.right_margin)

    for para in doc.paragraphs:
        if not para.runs:
            continue
        has_image = False
        for run in para.runs:
            if run._r.findall('.//' + qn('a:blip'), run._r.nsmap if hasattr(run._r, 'nsmap') else {}):
                has_image = True
                break
            # Check via inline shape tag
            if run._r.find(qn('w:drawing')) is not None:
                has_image = True
                break

        if not has_image:
            # Also check via inline shapes directly
            pass

    # Use InlineShape API — always force every image to full text width
    for shape in doc.inline_shapes:
        orig_w = shape.width
        orig_h = shape.height
        if orig_w is None or orig_h is None or orig_w == 0:
            continue
        full_w = int(text_width_emu)
        ratio  = full_w / orig_w
        shape.width  = full_w
        shape.height = int(orig_h * ratio)

    # Center-align paragraphs that contain images and remove first-line indent
    for para in doc.paragraphs:
        has_drawing = False
        for run in para.runs:
            if run._r.find(qn('w:drawing')) is not None:
                has_drawing = True
                break
        if has_drawing:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python format_docx_arg.py <input> <output>")
        sys.exit(1)
    format_document(sys.argv[1], sys.argv[2])
    print("Formatting complete.")
